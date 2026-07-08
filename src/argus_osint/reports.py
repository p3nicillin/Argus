from __future__ import annotations

import csv
import html
import json
from pathlib import Path
from typing import Any

from .repository import Repository, now
from .security import SecurityBriefBuilder


class ReportEngine:
    def __init__(self, repository: Repository) -> None:
        self.repository = repository

    def snapshot(self, case_id: int) -> dict[str, Any]:
        return {
            "generated_at": now(),
            "investigation": self.repository.investigation(case_id),
            **{
                name: self.repository.rows(name, case_id)
                for name in (
                    "notes",
                    "entities",
                    "relationships",
                    "evidence",
                    "bookmarks",
                    "timeline_events",
                    "intelligence",
                    "audit_log",
                    "source_records",
                    "comments",
                    "locations",
                    "collection_jobs",
                    "correlation_suggestions",
                )
            },
        }

    def export(self, case_id: int, destination: Path, format_name: str | None = None) -> Path:
        format_name = (format_name or destination.suffix.lstrip(".")).lower()
        if format_name == "json":
            destination.write_text(
                json.dumps(self.snapshot(case_id), indent=2, ensure_ascii=False), encoding="utf-8"
            )
        elif format_name == "csv":
            self._csv(case_id, destination)
        elif format_name in {"md", "markdown"}:
            destination.write_text(self._markdown(case_id), encoding="utf-8")
        elif format_name in {"txt", "text"}:
            destination.write_text(self._plain(case_id), encoding="utf-8")
        elif format_name in {"html", "htm"}:
            destination.write_text(self._html(case_id), encoding="utf-8")
        elif format_name == "pdf":
            self._pdf(case_id, destination)
        elif format_name == "docx":
            self._docx(case_id, destination)
        else:
            raise ValueError(f"Unsupported report format: {format_name}")
        return destination

    def export_security_brief(
        self,
        case_id: int,
        destination: Path,
        format_name: str | None = None,
    ) -> Path:
        format_name = (format_name or destination.suffix.lstrip(".")).lower()
        builder = SecurityBriefBuilder(self.repository)
        if format_name == "json":
            destination.write_text(
                json.dumps(builder.build(case_id), indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
        elif format_name in {"md", "markdown"}:
            destination.write_text(builder.markdown(case_id), encoding="utf-8")
        elif format_name in {"txt", "text"}:
            destination.write_text(
                builder.markdown(case_id).replace("#", "").replace("`", ""),
                encoding="utf-8",
            )
        elif format_name in {"html", "htm"}:
            brief = builder.build(case_id)
            risks = "".join(
                "<article>"
                f"<h3>{html.escape(str(item['score']))} {html.escape(item['title'])}</h3>"
                f"<p>{html.escape(item['kind'])}: {html.escape(item['value'])}</p>"
                "<ul>"
                + "".join(f"<li>{html.escape(reason)}</li>" for reason in item["reasons"])
                + "</ul></article>"
                for item in brief["top_risks"]
            )
            recommendations = "".join(
                f"<li><code>{html.escape(item['collector'])}</code> on "
                f"<code>{html.escape(item['query'])}</code>: {html.escape(item['reason'])}</li>"
                for item in brief["recommended_collection"]
            )
            destination.write_text(
                f"""<!doctype html><html lang=en><head><meta charset=utf-8><title>{html.escape(brief["investigation"]["title"])} security brief</title><style>
body{{font:14px system-ui,sans-serif;color:#18202b;max-width:960px;margin:40px auto;padding:0 28px;background:#f4f7fa}}header,article,section{{background:white;border:1px solid #dce3ea;border-radius:8px;padding:18px;margin:12px 0}}h1{{margin:0}}h2{{color:#9b2727;border-bottom:2px solid #d25a45;padding-bottom:6px}}code{{background:#eef2f6;padding:2px 4px;border-radius:4px}}.meta{{color:#607080}}
</style></head><body><header><h1>{html.escape(brief["investigation"]["title"])}</h1><p class=meta>Security research brief · Generated: {now()}</p></header><section><h2>Summary</h2><p>{brief["summary"]["risk_count"]} risks · {brief["summary"]["entity_count"]} entities · {brief["summary"]["source_count"]} sources</p></section><section><h2>Top risks</h2>{risks or '<p>No prioritized security risks were found.</p>'}</section><section><h2>Recommended collection</h2><ul>{recommendations}</ul></section></body></html>""",
                encoding="utf-8",
            )
        else:
            raise ValueError(f"Unsupported security brief format: {format_name}")
        return destination

    def _sections(self, case_id: int) -> list[tuple[str, list[dict[str, Any]]]]:
        return [
            (title, self.repository.rows(table, case_id))
            for title, table in (
                ("Entities", "entities"),
                ("Relationships", "relationships"),
                ("Evidence", "evidence"),
                ("Timeline", "timeline_events"),
                ("Intelligence", "intelligence"),
                ("Notes", "notes"),
                ("Bookmarks", "bookmarks"),
                ("Source provenance", "source_records"),
                ("Locations", "locations"),
                ("Investigator comments", "comments"),
                ("Collection jobs", "collection_jobs"),
                ("Correlation review", "correlation_suggestions"),
                ("Audit trail", "audit_log"),
            )
        ]

    def _markdown(self, case_id: int) -> str:
        case = self.repository.investigation(case_id)
        lines = [
            f"# {case['title']}",
            "",
            case["description"],
            "",
            f"**Status:** {case['status']}  ",
            f"**Investigator:** {case['investigator']}  ",
            f"**Generated:** {now()}",
            "",
        ]
        for title, rows in self._sections(case_id):
            lines.extend((f"## {title}", ""))
            if not rows:
                lines.extend(("_None recorded._", ""))
                continue
            for row in rows:
                heading = (
                    row.get("title")
                    or row.get("display_name")
                    or row.get("value")
                    or f"Record {row['id']}"
                )
                lines.append(f"### {heading}")
                for key, value in row.items():
                    if key not in {
                        "id",
                        "investigation_id",
                        "title",
                        "display_name",
                    } and value not in ("", None, [], {}):
                        rendered = (
                            json.dumps(value, ensure_ascii=False)
                            if isinstance(value, (dict, list))
                            else str(value)
                        )
                        lines.append(f"- **{key.replace('_', ' ').title()}:** {rendered}")
                lines.append("")
        return "\n".join(lines)

    def _plain(self, case_id: int) -> str:
        markdown = self._markdown(case_id)
        return (
            markdown.replace("### ", "")
            .replace("## ", "")
            .replace("# ", "")
            .replace("**", "")
            .replace("_None recorded._", "None recorded.")
        )

    def _html(self, case_id: int) -> str:
        case = self.repository.investigation(case_id)
        sections = []
        for title, rows in self._sections(case_id):
            cards = []
            for row in rows:
                heading = (
                    row.get("title")
                    or row.get("display_name")
                    or row.get("value")
                    or f"Record {row['id']}"
                )
                details = "".join(
                    f"<dt>{html.escape(key.replace('_', ' ').title())}</dt><dd>{html.escape(json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value))}</dd>"
                    for key, value in row.items()
                    if key not in {"id", "investigation_id", "title", "display_name"}
                    and value not in ("", None, [], {})
                )
                cards.append(
                    f"<article><h3>{html.escape(str(heading))}</h3><dl>{details}</dl></article>"
                )
            sections.append(
                f"<section><h2>{title}</h2>{''.join(cards) or '<p class=empty>None recorded.</p>'}</section>"
            )
        return f"""<!doctype html><html lang=en><head><meta charset=utf-8><title>{html.escape(case["title"])}</title><style>
body{{font:14px system-ui,sans-serif;color:#18202b;max-width:1100px;margin:40px auto;padding:0 28px;background:#f4f7fa}}header,article{{background:white;border:1px solid #dce3ea;border-radius:8px;padding:18px;margin:12px 0}}h1{{margin:0}}h2{{color:#196f9e;border-bottom:2px solid #45a4d1;padding-bottom:6px}}h3{{margin-top:0}}dl{{display:grid;grid-template-columns:180px 1fr;gap:6px 14px}}dt{{font-weight:600}}dd{{margin:0;overflow-wrap:anywhere}}.meta,.empty{{color:#607080}}@media print{{body{{background:white;margin:0}}article,header{{break-inside:avoid}}}}
</style></head><body><header><h1>{html.escape(case["title"])}</h1><p>{html.escape(case["description"])}</p><p class=meta>Status: {case["status"]} · Investigator: {html.escape(case["investigator"])} · Generated: {now()}</p></header>{"".join(sections)}</body></html>"""

    def _csv(self, case_id: int, destination: Path) -> None:
        data = self.snapshot(case_id)
        with destination.open("w", newline="", encoding="utf-8-sig") as stream:
            writer = csv.writer(stream)
            writer.writerow(("section", "id", "field", "value"))
            for section, records in data.items():
                if not isinstance(records, list):
                    continue
                for record in records:
                    for field, value in record.items():
                        writer.writerow(
                            (
                                section,
                                record.get("id", ""),
                                field,
                                json.dumps(value, ensure_ascii=False)
                                if isinstance(value, (list, dict))
                                else value,
                            )
                        )

    def _pdf(self, case_id: int, destination: Path) -> None:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError as exc:
            raise RuntimeError("PDF export requires reportlab") from exc
        case = self.repository.investigation(case_id)
        styles = getSampleStyleSheet()
        document = SimpleDocTemplate(
            str(destination),
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=case["title"],
            author=case["investigator"],
        )
        story = [
            Paragraph(html.escape(case["title"]), styles["Title"]),
            Paragraph(html.escape(case["description"]), styles["BodyText"]),
            Spacer(1, 8),
        ]
        for title, rows in self._sections(case_id):
            story.append(Paragraph(title, styles["Heading1"]))
            if not rows:
                story.append(Paragraph("None recorded.", styles["Italic"]))
            for row in rows:
                heading = (
                    row.get("title")
                    or row.get("display_name")
                    or row.get("value")
                    or f"Record {row['id']}"
                )
                story.append(Paragraph(html.escape(str(heading)), styles["Heading3"]))
                for key, value in row.items():
                    if key not in {
                        "id",
                        "investigation_id",
                        "title",
                        "display_name",
                    } and value not in ("", None, [], {}):
                        rendered = (
                            json.dumps(value, ensure_ascii=False)
                            if isinstance(value, (dict, list))
                            else str(value)
                        )
                        story.append(
                            Paragraph(
                                f"<b>{html.escape(key.replace('_', ' ').title())}:</b> {html.escape(rendered)}",
                                styles["BodyText"],
                            )
                        )
                story.append(Spacer(1, 5))
        document.build(story)

    def _docx(self, case_id: int, destination: Path) -> None:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX export requires python-docx") from exc
        case = self.repository.investigation(case_id)
        document = Document()
        document.core_properties.title = case["title"]
        document.core_properties.author = case["investigator"]
        document.add_heading(case["title"], 0)
        document.add_paragraph(case["description"])
        document.add_paragraph(
            f"Status: {case['status']} | Investigator: {case['investigator']} | Generated: {now()}"
        )
        for title, rows in self._sections(case_id):
            document.add_heading(title, level=1)
            if not rows:
                document.add_paragraph("None recorded.")
            for row in rows:
                heading = (
                    row.get("title")
                    or row.get("display_name")
                    or row.get("value")
                    or f"Record {row['id']}"
                )
                document.add_heading(str(heading), level=2)
                for key, value in row.items():
                    if key not in {
                        "id",
                        "investigation_id",
                        "title",
                        "display_name",
                    } and value not in ("", None, [], {}):
                        document.add_paragraph(
                            f"{key.replace('_', ' ').title()}: {json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else value}"
                        )
        document.save(destination)
